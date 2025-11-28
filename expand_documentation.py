#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Разширяване на документацията с допълнителни секции за дипломен проект
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_bullet_point(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def expand_documentation(input_file):
    doc = Document(input_file)
    
    # Добавяне на нови секции в края
    doc.add_page_break()
    
    # Детайлна имплементация
    add_heading(doc, '12. Детайлна имплементация на ключови функции', 1)
    
    add_heading(doc, '12.1. Backend - Document Generator', 2)
    add_paragraph(doc, 'Файл: backend/core/utils/document_generator.py', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Основната функция за генериране на документи:')
    
    add_code_block(doc, '''def generate_document_from_template(template_name, context):
    """
    Генерира Word документ от шаблон с контекст данни
    
    Args:
        template_name: Име на шаблон файл (напр. 'act14_bg.docx')
        context: Речник с данни за заместване
        
    Returns:
        Път до генерирания документ
    """
    # 1. Зареждане на шаблон
    template_path = os.path.join(settings.MEDIA_ROOT, 'templates', template_name)
    doc = Document(template_path)
    
    # 2. Bidirectional mapping
    if 'consultant_name' in context:
        context['representative_supervision'] = context['consultant_name']
    if 'designer_name' in context:
        context['representative_designer'] = context['designer_name']
    if 'contractor_name' in context:
        context['representative_builder'] = context['contractor_name']
    
    # 3. Замяна в параграфи
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # 4. Замяна в таблици
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    # 5. Премахване на незапълнени placeholders
    import re
    pattern = r'\\{\\{[^}]+\\}\\}'
    for paragraph in doc.paragraphs:
        paragraph.text = re.sub(pattern, '', paragraph.text)
    
    # 6. Запазване на документ
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'{template_name.replace(".docx", "")}_{timestamp}.docx'
    output_path = os.path.join(settings.MEDIA_ROOT, 'generated', filename)
    doc.save(output_path)
    
    return output_path''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Ключови аспекти:')
    add_bullet_point(doc, 'Използва python-docx библиотека за манипулация на Word файлове')
    add_bullet_point(doc, 'Bidirectional mapping осигурява consistency между различни полета')
    add_bullet_point(doc, 'Regex pattern премахва всички незапълнени {{placeholder}} маркери')
    add_bullet_point(doc, 'Timestamp в име на файл гарантира уникалност')
    add_bullet_point(doc, 'Обработва както параграфи, така и таблици')
    
    add_heading(doc, '12.2. Frontend - API Client', 2)
    add_paragraph(doc, 'Файл: frontend/src/api/client.ts', bold=True)
    add_paragraph(doc, '')
    
    add_code_block(doc, '''import axios from 'axios';

// Базова конфигурация
export const api = axios.create({
  baseURL: 'http://localhost:8000/api/',
  withCredentials: true,
  headers: {
    'Content-Type': 'application/json',
  },
});

// Функция за задаване на auth token
export const setAuthHeader = (token: string | null) => {
  if (token) {
    api.defaults.headers.common['Authorization'] = `Bearer ${token}`;
  } else {
    delete api.defaults.headers.common['Authorization'];
  }
};

// Interceptor за error handling
api.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      // Logout on unauthorized
      localStorage.removeItem('token');
      window.location.href = '/login';
    }
    return Promise.reject(error);
  }
);''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Особености:')
    add_bullet_point(doc, 'withCredentials: true позволява изпращане на cookies')
    add_bullet_point(doc, 'Bearer token authentication')
    add_bullet_point(doc, 'Автоматичен redirect към /login при 401 грешка')
    add_bullet_point(doc, 'Централизирана error handling логика')
    
    add_heading(doc, '12.3. Redux State Management', 2)
    add_paragraph(doc, 'Файл: frontend/src/store/authSlice.ts', bold=True)
    add_paragraph(doc, '')
    
    add_code_block(doc, '''import { createSlice, PayloadAction } from '@reduxjs/toolkit';

interface AuthState {
  user: User | null;
  token: string | null;
  isAdmin: boolean;
  isAuthenticated: boolean;
}

const initialState: AuthState = {
  user: null,
  token: localStorage.getItem('token'),
  isAdmin: false,
  isAuthenticated: false,
};

export const authSlice = createSlice({
  name: 'auth',
  initialState,
  reducers: {
    setUser: (state, action: PayloadAction<{user: User; token: string}>) => {
      state.user = action.payload.user;
      state.token = action.payload.token;
      state.isAdmin = action.payload.user.is_admin;
      state.isAuthenticated = true;
      localStorage.setItem('token', action.payload.token);
    },
    logout: (state) => {
      state.user = null;
      state.token = null;
      state.isAdmin = false;
      state.isAuthenticated = false;
      localStorage.removeItem('token');
    },
  },
});

export const { setUser, logout } = authSlice.actions;
export default authSlice.reducer;''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Redux Toolkit предимства:')
    add_bullet_point(doc, 'По-малко boilerplate код от класически Redux')
    add_bullet_point(doc, 'Вградена Immer библиотека за immutable updates')
    add_bullet_point(doc, 'TypeScript типизация out-of-the-box')
    add_bullet_point(doc, 'Автоматично генериране на action creators')
    
    doc.add_page_break()
    
    # Сигурност
    add_heading(doc, '13. Сигурност и оторизация', 1)
    
    add_heading(doc, '13.1. Django Permissions', 2)
    add_paragraph(doc, 'Файл: backend/core/permissions.py', bold=True)
    add_paragraph(doc, '')
    
    add_code_block(doc, '''from rest_framework import permissions

class IsEmployeeOrAdmin(permissions.BasePermission):
    """
    Достъп само за служители и администратори
    """
    def has_permission(self, request, view):
        return (
            request.user and 
            request.user.is_authenticated and 
            hasattr(request.user, 'employee')
        )

class IsAdminUser(permissions.BasePermission):
    """
    Достъп само за администратори
    """
    def has_permission(self, request, view):
        return (
            request.user and 
            request.user.is_authenticated and 
            hasattr(request.user, 'employee') and
            request.user.employee.is_admin
        )''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Използване в views:')
    add_code_block(doc, '''@api_view(['POST'])
@permission_classes([IsEmployeeOrAdmin])
def generate_document_view(request):
    # Само служители могат да генерират документи
    ...

@api_view(['DELETE'])
@permission_classes([IsAdminUser])
def delete_project_view(request, pk):
    # Само администратори могат да изтриват проекти
    ...''')
    
    add_heading(doc, '13.2. CORS Configuration', 2)
    add_paragraph(doc, 'Cross-Origin Resource Sharing настройки:')
    add_paragraph(doc, '')
    
    add_code_block(doc, '''# settings.py
INSTALLED_APPS = [
    ...
    'corsheaders',
]

MIDDLEWARE = [
    'corsheaders.middleware.CorsMiddleware',
    ...
]

CORS_ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://localhost:5174",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:5174",
]

CORS_ALLOW_CREDENTIALS = True''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Защо е важно:')
    add_bullet_point(doc, 'Frontend и backend са на различни портове (5173/5174 vs 8000)')
    add_bullet_point(doc, 'Браузърите блокират cross-origin requests по подразбиране')
    add_bullet_point(doc, 'CORS_ALLOW_CREDENTIALS позволява cookies и auth headers')
    
    add_heading(doc, '13.3. Password Security', 2)
    add_bullet_point(doc, 'Django използва PBKDF2 algorithm с SHA256 hash')
    add_bullet_point(doc, 'Минимум 100,000 iterations за хеширане')
    add_bullet_point(doc, 'Automatic salting на всяка парола')
    add_bullet_point(doc, 'Password validation rules (минимална дължина, complexity)')
    
    doc.add_page_break()
    
    # Тестване
    add_heading(doc, '14. Тестване на системата', 1)
    
    add_heading(doc, '14.1. Unit Tests', 2)
    add_paragraph(doc, 'Пример за Django test:')
    add_paragraph(doc, '')
    
    add_code_block(doc, '''from django.test import TestCase
from core.models import Project, Client

class ProjectModelTest(TestCase):
    def setUp(self):
        self.client = Client.objects.create(
            name="Тест Клиент",
            email="test@example.com"
        )
    
    def test_create_project(self):
        project = Project.objects.create(
            name="Тестов проект",
            location="гр. София",
            client=self.client,
            status="active",
            progress=0
        )
        self.assertEqual(project.name, "Тестов проект")
        self.assertEqual(project.status, "active")
        self.assertEqual(project.progress, 0)
    
    def test_project_progress_validation(self):
        project = Project.objects.create(
            name="Проект 2",
            client=self.client,
            progress=150  # Invalid
        )
        # Should not exceed 100
        self.assertLessEqual(project.progress, 100)''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Изпълнение на тестове:')
    add_code_block(doc, '''python manage.py test
python manage.py test core.tests.ProjectModelTest''')
    
    add_heading(doc, '14.2. Integration Tests', 2)
    add_paragraph(doc, 'API endpoint testing:')
    add_paragraph(doc, '')
    
    add_code_block(doc, '''from rest_framework.test import APITestCase
from rest_framework import status

class DocumentAPITest(APITestCase):
    def setUp(self):
        # Create test user and login
        self.user = User.objects.create_user(
            username='testuser',
            password='testpass123'
        )
        self.client.force_authenticate(user=self.user)
    
    def test_generate_document(self):
        data = {
            'template_name': 'act14_bg.docx',
            'context': {
                'project_name': 'Тест проект',
                'act_date': '28.11.2025'
            }
        }
        response = self.client.post('/api/documents/generate/', data, format='json')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('file_docx', response.data)''')
    
    add_heading(doc, '14.3. Frontend Testing', 2)
    add_paragraph(doc, 'Jest + React Testing Library:')
    add_paragraph(doc, '')
    
    add_code_block(doc, '''import { render, screen } from '@testing-library/react';
import Header from './Header';

test('renders header with logo', () => {
  render(<Header />);
  const logoElement = screen.getByText(/Строителен Надзор/i);
  expect(logoElement).toBeInTheDocument();
});''')
    
    doc.add_page_break()
    
    # Performance
    add_heading(doc, '15. Производителност и оптимизация', 1)
    
    add_heading(doc, '15.1. Backend оптимизации', 2)
    add_bullet_point(doc, 'Database indexing на често търсени полета')
    add_bullet_point(doc, 'QuerySet optimization с select_related() и prefetch_related()')
    add_bullet_point(doc, 'Caching на статични данни с Django cache framework')
    add_bullet_point(doc, 'Pagination на API endpoints за големи datasets')
    add_bullet_point(doc, 'Compression на HTTP responses (gzip)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Пример за optimized query:')
    add_code_block(doc, '''# Неоптимизирано (N+1 problem)
projects = Project.objects.all()
for project in projects:
    print(project.client.name)  # Extra query за всеки project

# Оптимизирано
projects = Project.objects.select_related('client').all()
for project in projects:
    print(project.client.name)  # Само 1 query''')
    
    add_heading(doc, '15.2. Frontend оптимизации', 2)
    add_bullet_point(doc, 'Code splitting с React.lazy() и Suspense')
    add_bullet_point(doc, 'Memoization с useMemo и useCallback hooks')
    add_bullet_point(doc, 'Virtualization на дълги списъци')
    add_bullet_point(doc, 'Image optimization и lazy loading')
    add_bullet_point(doc, 'Bundle size optimization (tree shaking)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Пример за code splitting:')
    add_code_block(doc, '''import { lazy, Suspense } from 'react';

const Documents = lazy(() => import('./pages/Documents'));
const Projects = lazy(() => import('./pages/Projects'));

function App() {
  return (
    <Suspense fallback={<div>Loading...</div>}>
      <Routes>
        <Route path="/documents" element={<Documents />} />
        <Route path="/projects" element={<Projects />} />
      </Routes>
    </Suspense>
  );
}''')
    
    add_heading(doc, '15.3. Database оптимизация', 2)
    add_paragraph(doc, 'Индекси за бърз search:')
    add_code_block(doc, '''class Project(models.Model):
    name = models.CharField(max_length=200, db_index=True)
    status = models.CharField(max_length=20, db_index=True)
    
    class Meta:
        indexes = [
            models.Index(fields=['status', 'start_date']),
            models.Index(fields=['client', 'status']),
        ]''')
    
    doc.add_page_break()
    
    # Грешки и решения
    add_heading(doc, '16. Често срещани проблеми и решения', 1)
    
    add_heading(doc, '16.1. CORS Errors', 2)
    add_paragraph(doc, 'Проблем:', bold=True)
    add_paragraph(doc, 'Access to XMLHttpRequest has been blocked by CORS policy')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:', bold=True)
    add_bullet_point(doc, '1. Инсталирайте django-cors-headers')
    add_bullet_point(doc, '2. Добавете в INSTALLED_APPS и MIDDLEWARE')
    add_bullet_point(doc, '3. Конфигурирайте CORS_ALLOWED_ORIGINS')
    add_bullet_point(doc, '4. Задайте CORS_ALLOW_CREDENTIALS = True')
    
    add_heading(doc, '16.2. Template Not Found', 2)
    add_paragraph(doc, 'Проблем:', bold=True)
    add_paragraph(doc, 'FileNotFoundError: [Errno 2] No such file or directory: act14_bg.docx')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:', bold=True)
    add_bullet_point(doc, 'Проверете че шаблонът е в backend/media/templates/')
    add_bullet_point(doc, 'Проверете правописа на името на файла')
    add_bullet_point(doc, 'Уверете се че MEDIA_ROOT е правилно конфигуриран')
    
    add_heading(doc, '16.3. Authentication Failed', 2)
    add_paragraph(doc, 'Проблем:', bold=True)
    add_paragraph(doc, 'HTTP 401 Unauthorized при API заявки')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:', bold=True)
    add_bullet_point(doc, 'Проверете че token е съхранен в localStorage')
    add_bullet_point(doc, 'Уверете се че setAuthHeader() е извикана след login')
    add_bullet_point(doc, 'Проверете Authorization header във Network tab')
    add_bullet_point(doc, 'Проверете че token не е expired')
    
    add_heading(doc, '16.4. Port Already in Use', 2)
    add_paragraph(doc, 'Проблем:', bold=True)
    add_paragraph(doc, 'Error: That port is already in use')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение за Windows:', bold=True)
    add_code_block(doc, '''# Намиране на процес на порт 8000
netstat -ano | findstr :8000

# Убиване на процес (заменете PID)
taskkill /PID <process_id> /F''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Или променете порта:')
    add_code_block(doc, '''python manage.py runserver 8001
npm run dev -- --port 5175''')
    
    doc.add_page_break()
    
    # Бъдещо развитие
    add_heading(doc, '17. Бъдещо развитие и подобрения', 1)
    
    add_heading(doc, '17.1. Планирани функционалности', 2)
    
    add_paragraph(doc, 'Фаза 1 (Краткосрочни - 1-3 месеца):', bold=True)
    add_bullet_point(doc, 'PDF генериране директно от backend (без LibreOffice)')
    add_bullet_point(doc, 'Batch генериране на множество актове')
    add_bullet_point(doc, 'Email notifications за важни събития')
    add_bullet_point(doc, 'Export на проекти в Excel/CSV')
    add_bullet_point(doc, 'Advanced search и filtering')
    add_bullet_point(doc, 'Document versioning (revision history)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Фаза 2 (Средносрочни - 3-6 месеца):', bold=True)
    add_bullet_point(doc, 'Електронни подписи на документи (qualified electronic signature)')
    add_bullet_point(doc, 'OCR за сканирани документи')
    add_bullet_point(doc, 'Gantt chart за проектно планиране')
    add_bullet_point(doc, 'Dashboard с analytics и reports')
    add_bullet_point(doc, 'Mobile app (React Native)')
    add_bullet_point(doc, 'Integration с външни системи (ERP, CAD)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Фаза 3 (Дългосрочни - 6-12 месеца):', bold=True)
    add_bullet_point(doc, 'AI-powered document analysis')
    add_bullet_point(doc, 'Automated compliance checking')
    add_bullet_point(doc, 'Real-time collaboration (WebSocket)')
    add_bullet_point(doc, 'BIM (Building Information Modeling) integration')
    add_bullet_point(doc, 'Multi-language support')
    add_bullet_point(doc, 'Blockchain за document provenance')
    
    add_heading(doc, '17.2. Технологични подобрения', 2)
    add_bullet_point(doc, 'Миграция към PostgreSQL за production')
    add_bullet_point(doc, 'Redis за caching и session storage')
    add_bullet_point(doc, 'Elasticsearch за full-text search')
    add_bullet_point(doc, 'GraphQL API като алтернатива на REST')
    add_bullet_point(doc, 'Docker containerization')
    add_bullet_point(doc, 'Kubernetes orchestration за scaling')
    add_bullet_point(doc, 'CI/CD pipeline с GitHub Actions')
    
    add_heading(doc, '17.3. UX подобрения', 2)
    add_bullet_point(doc, 'Drag-and-drop file upload')
    add_bullet_point(doc, 'Inline editing в таблици')
    add_bullet_point(doc, 'Keyboard shortcuts')
    add_bullet_point(doc, 'Dark mode (вече имплементирано частично)')
    add_bullet_point(doc, 'Customizable dashboard widgets')
    add_bullet_point(doc, 'Tour guide за нови потребители')
    
    doc.add_page_break()
    
    # Заключение разширено
    add_heading(doc, '18. Технически спецификации', 1)
    
    add_heading(doc, '18.1. Системни изисквания', 2)
    
    add_paragraph(doc, 'Development:', bold=True)
    add_bullet_point(doc, 'OS: Windows 10/11, macOS 10.15+, Ubuntu 20.04+')
    add_bullet_point(doc, 'RAM: Минимум 4GB (препоръчва се 8GB)')
    add_bullet_point(doc, 'Disk: 2GB свободно пространство')
    add_bullet_point(doc, 'CPU: Dual-core 2.0GHz+')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Production Server:', bold=True)
    add_bullet_point(doc, 'RAM: Минимум 2GB (препоръчва се 4GB)')
    add_bullet_point(doc, 'Disk: 10GB+ (зависи от обем документи)')
    add_bullet_point(doc, 'CPU: 2+ vCPUs')
    add_bullet_point(doc, 'Bandwidth: 100Mbps+')
    
    add_heading(doc, '18.2. Browser Support', 2)
    add_bullet_point(doc, 'Chrome 90+ (препоръчва се)')
    add_bullet_point(doc, 'Firefox 88+')
    add_bullet_point(doc, 'Edge 90+')
    add_bullet_point(doc, 'Safari 14+')
    add_bullet_point(doc, 'Opera 76+')
    
    add_heading(doc, '18.3. Мрежови изисквания', 2)
    add_bullet_point(doc, 'HTTP/HTTPS протокол')
    add_bullet_point(doc, 'WebSocket support (за бъдещи real-time features)')
    add_bullet_point(doc, 'Порт 8000 (backend) и 5173/5174 (frontend) отворени')
    add_bullet_point(doc, 'Firewall rules за production deployment')
    
    doc.add_page_break()
    
    # Речник на термини
    add_heading(doc, '19. Речник на термини', 1)
    
    add_paragraph(doc, 'API (Application Programming Interface):', bold=True)
    add_paragraph(doc, 'Набор от функции и процедури, които позволяват създаването на приложения, които имат достъп до функции на операционна система, приложение или друга услуга.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'REST (Representational State Transfer):', bold=True)
    add_paragraph(doc, 'Архитектурен стил за проектиране на мрежови приложения, използващ HTTP методи.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'JWT (JSON Web Token):', bold=True)
    add_paragraph(doc, 'Компактен, URL-safe начин за представяне на claims за трансфер между две страни.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'CORS (Cross-Origin Resource Sharing):', bold=True)
    add_paragraph(doc, 'Механизъм, който използва допълнителни HTTP headers за да каже на браузъра да даде разрешение на уеб приложение да има достъп до избрани ресурси от различен origin.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'SPA (Single Page Application):', bold=True)
    add_paragraph(doc, 'Уеб приложение, което взаимодейства с потребителя чрез динамично пренаписване на текущата страница вместо зареждане на цели нови страници от сървър.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'ORM (Object-Relational Mapping):', bold=True)
    add_paragraph(doc, 'Техника за конвертиране на данни между несъвместими типови системи чрез обектно-ориентиран програмен език.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'Migration:', bold=True)
    add_paragraph(doc, 'Версионен контрол за database schema. Позволява да се променя структурата на базата данни по контролиран начин.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'Serializer:', bold=True)
    add_paragraph(doc, 'Компонент, който преобразува сложни типове данни (model instances) в Python data types, които след това могат лесно да бъдат рендирани в JSON.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'ViewSet:', bold=True)
    add_paragraph(doc, 'Django REST Framework клас, който комбинира логиката за множество свързани views в един клас.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'Middleware:', bold=True)
    add_paragraph(doc, 'Софтуер, който действа като мост между операционна система и приложения, или между frontend и backend.')
    
    doc.add_page_break()
    
    # Библиография
    add_heading(doc, '20. Използвани библиотеки и ресурси', 1)
    
    add_heading(doc, '20.1. Backend Dependencies', 2)
    add_paragraph(doc, 'requirements.txt:', italic=True)
    add_code_block(doc, '''Django==5.2.8
djangorestframework==3.15.2
django-cors-headers==4.6.0
python-docx==1.2.0
python-dotenv==1.0.0
Pillow==11.0.0
PyJWT==2.10.1
gunicorn==23.0.0''')
    
    add_heading(doc, '20.2. Frontend Dependencies', 2)
    add_paragraph(doc, 'package.json:', italic=True)
    add_code_block(doc, '''{
  "dependencies": {
    "react": "^18.3.1",
    "react-dom": "^18.3.1",
    "react-router-dom": "^7.1.1",
    "@reduxjs/toolkit": "^2.5.0",
    "react-redux": "^9.2.0",
    "axios": "^1.7.9",
    "antd": "^5.22.5",
    "dayjs": "^1.11.13",
    "styled-components": "^6.1.13",
    "framer-motion": "^12.23.24"
  },
  "devDependencies": {
    "typescript": "^5.6.3",
    "vite": "^6.0.3",
    "@vitejs/plugin-react": "^4.3.4"
  }
}''')
    
    add_heading(doc, '20.3. Полезни ресурси', 2)
    add_bullet_point(doc, 'Django Documentation: https://docs.djangoproject.com/')
    add_bullet_point(doc, 'Django REST Framework: https://www.django-rest-framework.org/')
    add_bullet_point(doc, 'React Documentation: https://react.dev/')
    add_bullet_point(doc, 'TypeScript Handbook: https://www.typescriptlang.org/docs/')
    add_bullet_point(doc, 'Ant Design: https://ant.design/')
    add_bullet_point(doc, 'python-docx Documentation: https://python-docx.readthedocs.io/')
    add_bullet_point(doc, 'MDN Web Docs: https://developer.mozilla.org/')
    
    doc.add_page_break()
    
    # Благодарности и Автор
    add_heading(doc, 'Заключителни думи', 1)
    
    add_paragraph(doc, 'Системата за Строителен Надзор представлява модерно, скалируемо и сигурно решение за управление на строителни проекти и автоматизация на документацията. Проектът демонстрира best practices в уеб разработката и може да служи като основа за бъдещо разширяване.')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'Ключови постижения:', bold=True)
    add_bullet_point(doc, 'Пълна интеграция между backend (Django) и frontend (React)')
    add_bullet_point(doc, 'Автоматизация на генериране на официални документи')
    add_bullet_point(doc, 'Сигурна authentication и authorization система')
    add_bullet_point(doc, 'Responsive дизайн за всички устройства')
    add_bullet_point(doc, 'Чист, поддържаем код с TypeScript типизация')
    add_bullet_point(doc, 'Готовност за production deployment')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Системата е тествана и готова за реална употреба в строителни фирми, консултантски компании и организации, занимаващи се със строителен надзор.')
    
    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, '---', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Дата: Ноември 2025', italic=True)
    add_paragraph(doc, 'Версия: 1.0', italic=True)
    add_paragraph(doc, 'Статус: Завършен дипломен проект', italic=True)
    
    # Запазване
    output_file = 'Документация_Система_за_Строителен_Надзор_РАЗШИРЕНА.docx'
    doc.save(output_file)
    
    print(f'✓ Разширената документация е създадена: {output_file}')
    print(f'✓ Общо параграфи: {len(doc.paragraphs)}')
    print(f'✓ Приблизителен брой страници: {len(doc.paragraphs) // 25}')
    
    return output_file

if __name__ == '__main__':
    input_file = 'Документация_Система_за_Строителен_Надзор.docx'
    if os.path.exists(input_file):
        expand_documentation(input_file)
    else:
        print(f'Грешка: Файлът {input_file} не е намерен!')
