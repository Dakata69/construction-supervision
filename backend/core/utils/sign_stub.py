from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
from docx.shared import Inches
import os
from datetime import datetime

def create_signature(name, role, date=None):
    """
    Create a signature image with name, role, and date.
    
    Args:
        name (str): Name of the signer
        role (str): Role of the signer
        date (str): Date of signing (defaults to current date)
    
    Returns:
        BytesIO: Signature image in memory buffer
    """
    width = 600
    height = 200
    image = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(image)
    
    try:
        font = ImageFont.truetype("arial.ttf", 32)
    except:
        font = ImageFont.load_default()
    
    line_y = height - 60
    draw.line([(50, line_y), (width-50, line_y)], fill='black', width=2)
    
    draw.text((50, line_y - 40), name, fill='black', font=font)
    draw.text((50, line_y + 10), role, fill='black', font=font)
    
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")
    draw.text((width-150, line_y + 10), date, fill='black', font=font)
    
    img_buffer = BytesIO()
    image.save(img_buffer, format='PNG')
    img_buffer.seek(0)
    
    return img_buffer

def sign_document(doc, signatures):
    """
    Add signatures to a document.
    
    Args:
        doc: python-docx Document object
        signatures: List of tuples (name, role, date)
    """
    doc.add_heading('Signatures', level=1)
    
    table = doc.add_table(rows=1, cols=len(signatures))
    
    for idx, (name, role, date) in enumerate(signatures):
        sig_image = create_signature(name, role, date)
        
        cell = table.cell(0, idx)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(sig_image, width=Inches(2.5))
    
    return doc
