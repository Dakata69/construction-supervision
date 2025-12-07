import os
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
from .sign_stub import sign_document

logger = logging.getLogger(__name__)

TEMPLATE_DIR = os.path.join(settings.MEDIA_ROOT, 'templates')

def ensure_templates_dir():
    """Ensure the templates directory exists"""
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    return TEMPLATE_DIR

def get_template_path(template_name):
    """Get the full path for a template"""
    return os.path.join(TEMPLATE_DIR, template_name)

def generate_document(template_name, context, output_path, signatures=None):
    """
    Generate a document from a template and context.
    
    Args:
        template_name (str): Name of the template file (e.g. 'daily_report.docx')
        context (dict): Context data to fill in the template
        output_path (str): Path where to save the generated document
        signatures (dict, optional): Dictionary of signature placeholders and their values
    """
    template_path = get_template_path(template_name)
    logger.info(f'Using template at {template_path}')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'Template {template_name} not found at {template_path}')
    
    try:
        doc = Document(template_path)
        
        enriched_context = dict(context)
        
        if 'consultant_name' not in enriched_context and 'representative_supervision' in enriched_context:
            enriched_context['consultant_name'] = enriched_context['representative_supervision']
        if 'representative_supervision' not in enriched_context and 'consultant_name' in enriched_context:
            enriched_context['representative_supervision'] = enriched_context['consultant_name']
        
        if 'designer_name' not in enriched_context and 'representative_designer' in enriched_context:
            enriched_context['designer_name'] = enriched_context['representative_designer']
        if 'representative_designer' not in enriched_context and 'designer_name' in enriched_context:
            enriched_context['representative_designer'] = enriched_context['designer_name']
        
        if 'contractor_name' not in enriched_context and 'representative_builder' in enriched_context:
            enriched_context['contractor_name'] = enriched_context['representative_builder']
        if 'representative_builder' not in enriched_context and 'contractor_name' in enriched_context:
            enriched_context['representative_builder'] = enriched_context['contractor_name']
        
        if 'supervisor_name' in enriched_context:
            if 'consultant_name' not in enriched_context:
                enriched_context['consultant_name'] = enriched_context['supervisor_name']
            if 'representative_supervision' not in enriched_context:
                enriched_context['representative_supervision'] = enriched_context['supervisor_name']
        
        for paragraph in doc.paragraphs:
            for key, value in enriched_context.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in paragraph.text:
                    if value is not None and str(value).strip():
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                    else:
                        paragraph.text = paragraph.text.replace(placeholder, '')
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in enriched_context.items():
                            placeholder = f'{{{{{key}}}}}'
                            if placeholder in paragraph.text:
                                if value is not None and str(value).strip():
                                    paragraph.text = paragraph.text.replace(placeholder, str(value))
                                else:
                                    paragraph.text = paragraph.text.replace(placeholder, '')
        
        numeric_map = {
            '1': enriched_context.get('representative_builder') or enriched_context.get('contractor_name') or '',
            '2': enriched_context.get('representative_supervision') or enriched_context.get('consultant_name') or '',
            '3': enriched_context.get('representative_designer') or enriched_context.get('designer_name') or '',
        }
        def replace_numeric_markers(text: str) -> str:
            for num, val in numeric_map.items():
                marker = f'*{num}*'
                if marker in text and val:
                    text = text.replace(marker, str(val))
            return text
        
        for paragraph in doc.paragraphs:
            new_text = replace_numeric_markers(paragraph.text)
            if new_text != paragraph.text:
                paragraph.text = new_text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        new_text = replace_numeric_markers(paragraph.text)
                        if new_text != paragraph.text:
                            paragraph.text = new_text
        
        import re
        placeholder_pattern = re.compile(r'\{\{[^}]+\}\}')
        
        for paragraph in doc.paragraphs:
            if '{{' in paragraph.text:
                paragraph.text = placeholder_pattern.sub('', paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{' in paragraph.text:
                            paragraph.text = placeholder_pattern.sub('', paragraph.text)
        
        if signatures:
            doc = sign_document(doc, signatures)

        if template_name in ['act7_bg.docx', 'act14_bg.docx', 'act15_bg.docx']:
            pass
        else:
            if context.get('quality_control'):
                doc.add_heading('Quality Control', level=1)
                doc.add_paragraph(context['quality_control'])
            
            if context.get('issues'):
                doc.add_heading('Issues and Concerns', level=1)
                doc.add_paragraph(context['issues'])
            
            if context.get('materials_delivered'):
                doc.add_heading('Materials Delivered', level=1)
                doc.add_paragraph(context['materials_delivered'])
            
            if context.get('next_steps'):
                doc.add_heading('Next Steps', level=1)
                doc.add_paragraph(context['next_steps'])
            
            if context.get('notes'):
                doc.add_heading('Additional Notes', level=1)
                doc.add_paragraph(context['notes'])
                    
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        sigs = signatures or enriched_context.get('signatures')
        if sigs:
            doc = sign_document(doc, sigs)

        doc.save(output_path)
        logger.info(f'Document generated successfully at {output_path}')

    except Exception as e:
        logger.error(f'Document generation failed: {str(e)}')
        raise
