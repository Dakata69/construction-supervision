#!/usr/bin/env python3
"""
Convert markdown files from Desktop\DOCS folder to Word .docx format.
Preserves formatting: headings, bold, italics, code blocks, lists, tables.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Desktop DOCS folder
DESKTOP_DOCS_FOLDER = Path.home() / "Desktop" / "DOCS"
OUTPUT_FOLDER = DESKTOP_DOCS_FOLDER / "converted"

# Create output folder if it doesn't exist
OUTPUT_FOLDER.mkdir(exist_ok=True)

def parse_markdown_to_docx(md_file_path, docx_file_path):
    """
    Convert a markdown file to a Word document.
    Handles:
    - Headings (# ## ### etc.)
    - Bold (**text**)
    - Italic (*text*)
    - Code blocks (```...```)
    - Inline code (`code`)
    - Lists (- * +)
    - Tables (|...)
    """
    
    doc = Document()
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split by lines
    lines = content.split('\n')
    
    in_code_block = False
    code_block_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # End code block
                in_code_block = False
                if code_block_lines:
                    # Add code block to document
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    for code_line in code_block_lines:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                code_block_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.strip().startswith('# '):
            heading_level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=heading_level)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle tables (basic support)
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) > 1:
                # Parse table
                rows = [re.split(r'\|', row.strip('| ')) for row in table_lines]
                # Skip separator row (second row)
                rows = [rows[0]] + rows[2:] if len(rows) > 2 else rows
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for i_row, row_data in enumerate(rows):
                        for i_col, cell_text in enumerate(row_data):
                            cell = table.rows[i_row].cells[i_col]
                            cell.text = cell_text.strip()
            continue
        
        # Handle lists
        if line.strip() and re.match(r'^[\-\*\+]\s', line.strip()):
            list_text = re.sub(r'^[\-\*\+]\s+', '', line.strip())
            p = doc.add_paragraph(list_text, style='List Bullet')
            # Apply inline formatting
            apply_inline_formatting(p, list_text)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            apply_inline_formatting(p, line.strip())
        else:
            # Empty line = paragraph break
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file_path)
    print(f"✓ Converted: {md_file_path.name} → {docx_file_path.name}")

def apply_inline_formatting(paragraph, text):
    """
    Apply inline formatting (bold, italic, code) to a paragraph.
    """
    # Pattern: **bold**, *italic*, `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(128, 0, 0)  # Dark red
        else:
            # Plain text
            paragraph.add_run(part)

def main():
    """Convert all markdown files in Desktop\DOCS to Word format."""
    
    if not DESKTOP_DOCS_FOLDER.exists():
        print(f"❌ Folder not found: {DESKTOP_DOCS_FOLDER}")
        print("Please ensure you have a DOCS folder on your Desktop.")
        return
    
    md_files = list(DESKTOP_DOCS_FOLDER.glob("*.md"))
    
    if not md_files:
        print(f"⚠ No .md files found in {DESKTOP_DOCS_FOLDER}")
        return
    
    print(f"Found {len(md_files)} markdown file(s). Converting...\n")
    
    for md_file in md_files:
        docx_file = OUTPUT_FOLDER / f"{md_file.stem}.docx"
        try:
            parse_markdown_to_docx(md_file, docx_file)
        except Exception as e:
            print(f"❌ Error converting {md_file.name}: {e}")
    
    print(f"\n✅ Done! Word documents saved to: {OUTPUT_FOLDER}")

if __name__ == "__main__":
    main()
