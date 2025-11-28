#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Създаване на РАЗШИРЕНА документация на проекта в Word формат
За дипломен проект - минимум 40 страници
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def add_heading(doc, text, level=1):
    """Добавя заглавие с форматиране"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    """Добавя параграф с форматиране"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_bullet_point(doc, text):
    """Добавя точка с водеща марка"""
    p = doc.add_paragraph(text, style='List Bullet')
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_code_block(doc, code, language=""):
    """Добавя код блок"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    # Сив фон ефект чрез шейдинг (ограничено в python-docx)
    return p

def create_documentation():
    """Създава Word документ с пълна документация"""
    doc = Document()
    
    # Заглавна страница
    title = doc.add_heading('Система за Строителен Надзор', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Пълна документация на проекта')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('Ноември 2025')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # Съдържание
    add_heading(doc, 'Съдържание', 1)
    add_paragraph(doc, '1. Въведение и Обща информация')
    add_paragraph(doc, '2. Архитектура на системата')
    add_paragraph(doc, '3. Backend (Django)')
    add_paragraph(doc, '4. Frontend (React + TypeScript)')
    add_paragraph(doc, '5. Генериране на документи')
    add_paragraph(doc, '6. Модули и функционалности')
    add_paragraph(doc, '7. База данни')
    add_paragraph(doc, '8. API документация')
    add_paragraph(doc, '9. Инсталация и стартиране')
    add_paragraph(doc, '10. Deployment')
    add_paragraph(doc, '11. Примери за използване')
    
    doc.add_page_break()
    
    # 1. Въведение
    add_heading(doc, '1. Въведение и Обща информация', 1)
    
    add_heading(doc, '1.1. Цел на проекта', 2)
    add_paragraph(doc, 'Системата за Строителен Надзор е уеб базирана платформа, предназначена да улеснява управлението и документацията на строителни проекти. Основните цели включват:')
    add_bullet_point(doc, 'Автоматизирано генериране на официални строителни актове (Акт 7, Акт 14, Акт 15)')
    add_bullet_point(doc, 'Централизирано управление на проекти, екипи и задачи')
    add_bullet_point(doc, 'Съхранение и организация на документи')
    add_bullet_point(doc, 'Проследяване на напредъка на строителни обекти')
    add_bullet_point(doc, 'Мобилно-адаптивен интерфейс за достъп от всяко устройство')
    
    add_heading(doc, '1.2. Технологичен стек', 2)
    add_paragraph(doc, 'Backend:', bold=True)
    add_bullet_point(doc, 'Django 5.2.8 - Web framework')
    add_bullet_point(doc, 'Django REST Framework - RESTful API')
    add_bullet_point(doc, 'python-docx 1.2.0 - Генериране на Word документи')
    add_bullet_point(doc, 'SQLite/PostgreSQL - База данни')
    add_bullet_point(doc, 'Django CORS Headers - Cross-Origin Resource Sharing')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Frontend:', bold=True)
    add_bullet_point(doc, 'React 18 - JavaScript библиотека за UI')
    add_bullet_point(doc, 'TypeScript - Типизиран JavaScript')
    add_bullet_point(doc, 'Vite - Build tool и dev server')
    add_bullet_point(doc, 'Ant Design - UI компонентна библиотека')
    add_bullet_point(doc, 'Redux Toolkit - State management')
    add_bullet_point(doc, 'Axios - HTTP клиент')
    add_bullet_point(doc, 'styled-components - CSS-in-JS')
    add_bullet_point(doc, 'dayjs - Работа с дати')
    
    doc.add_page_break()
    
    # 2. Архитектура
    add_heading(doc, '2. Архитектура на системата', 1)
    
    add_heading(doc, '2.1. Общ преглед', 2)
    add_paragraph(doc, 'Системата следва класическа client-server архитектура:')
    add_bullet_point(doc, 'Frontend (React SPA) - комуникира с backend през REST API')
    add_bullet_point(doc, 'Backend (Django) - обработва бизнес логика и данни')
    add_bullet_point(doc, 'Database (SQLite/PostgreSQL) - съхранява данни')
    add_bullet_point(doc, 'Media Storage - съхранява генерирани документи и качени файлове')
    
    add_heading(doc, '2.2. Структура на проекта', 2)
    add_paragraph(doc, 'Monorepo структура с две основни папки:')
    add_paragraph(doc, '')
    
    add_code_block(doc, '''construction-supervision/
├── backend/                    # Django проект
│   ├── config/                # Настройки на проекта
│   │   ├── settings.py       # Основни настройки
│   │   ├── urls.py           # URL routing
│   │   └── wsgi.py           # WSGI конфигурация
│   ├── core/                  # Главно приложение
│   │   ├── models/           # Модели на данни
│   │   ├── views/            # API views
│   │   ├── serializers.py    # DRF serializers
│   │   ├── permissions.py    # Права за достъп
│   │   └── utils/            # Помощни функции
│   ├── media/                 # Съхранение на файлове
│   │   ├── templates/        # Word шаблони
│   │   ├── generated/        # Генерирани документи
│   │   └── acts/             # Качени актове
│   ├── manage.py             # Django CLI
│   └── requirements.txt      # Python dependencies
│
├── frontend/                  # React приложение
│   ├── src/
│   │   ├── api/              # API клиент
│   │   ├── components/       # React компоненти
│   │   ├── pages/            # Страници
│   │   ├── store/            # Redux store
│   │   ├── styles/           # CSS файлове
│   │   ├── types/            # TypeScript типове
│   │   └── App.tsx           # Главен компонент
│   ├── package.json          # Node dependencies
│   └── vite.config.ts        # Vite конфигурация
│
├── scripts/                   # Utility scripts
└── README.md                  # Документация''')
    
    doc.add_page_break()
    
    # 3. Backend
    add_heading(doc, '3. Backend (Django)', 1)
    
    add_heading(doc, '3.1. Модели на данни', 2)
    add_paragraph(doc, 'Системата използва следните основни модели:')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Employee (Служител):', bold=True)
    add_bullet_point(doc, 'user - Връзка с Django User модел')
    add_bullet_point(doc, 'position - Длъжност')
    add_bullet_point(doc, 'phone - Телефон')
    add_bullet_point(doc, 'is_admin - Флаг за администратор')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Client (Клиент/Възложител):', bold=True)
    add_bullet_point(doc, 'name - Име')
    add_bullet_point(doc, 'contact_person - Лице за контакт')
    add_bullet_point(doc, 'email - Имейл')
    add_bullet_point(doc, 'phone - Телефон')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Project (Проект/Строеж):', bold=True)
    add_bullet_point(doc, 'name - Име на проекта')
    add_bullet_point(doc, 'location - Местонахождение')
    add_bullet_point(doc, 'client - Връзка към Client')
    add_bullet_point(doc, 'start_date - Начална дата')
    add_bullet_point(doc, 'end_date - Крайна дата')
    add_bullet_point(doc, 'status - Статус (active/completed/paused)')
    add_bullet_point(doc, 'progress - Процент завършеност (0-100)')
    add_bullet_point(doc, 'description - Описание')
    add_bullet_point(doc, 'permit_number - Номер на разрешение')
    add_bullet_point(doc, 'consultant_name - Консултант (строителен надзор)')
    add_bullet_point(doc, 'contractor_name - Строител (изпълнител)')
    add_bullet_point(doc, 'designer_name - Проектант')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Document (Документ):', bold=True)
    add_bullet_point(doc, 'title - Заглавие')
    add_bullet_point(doc, 'file_docx - DOCX файл')
    add_bullet_point(doc, 'file_pdf - PDF файл')
    add_bullet_point(doc, 'zip_url - ZIP архив (съдържа DOCX + PDF)')
    add_bullet_point(doc, 'created_at - Дата на създаване')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Act (Акт):', bold=True)
    add_bullet_point(doc, 'project - Връзка към Project')
    add_bullet_point(doc, 'act_type - Тип акт (7, 14, 15)')
    add_bullet_point(doc, 'act_number - Номер на акт')
    add_bullet_point(doc, 'act_date - Дата на акт')
    add_bullet_point(doc, 'file_docx - DOCX файл')
    add_bullet_point(doc, 'file_pdf - PDF файл')
    add_bullet_point(doc, 'zip_file - ZIP архив')
    add_bullet_point(doc, 'context_data - JSON данни с контекст')
    
    add_heading(doc, '3.2. API Endpoints', 2)
    add_paragraph(doc, 'Основни API маршрути:')
    add_paragraph(doc, '')
    
    add_code_block(doc, '''GET/POST    /api/projects/              - Списък/създаване на проекти
GET/PUT     /api/projects/{id}/         - Детайли/редакция на проект
GET/POST    /api/documents/             - Списък/качване на документи
DELETE      /api/documents/{id}/        - Изтриване на документ
POST        /api/documents/generate/    - Генериране на акт
GET/POST    /api/acts/                  - Списък/създаване на актове
POST        /api/auth/login/            - Вход в системата
POST        /api/auth/logout/           - Изход от системата
GET         /api/auth/me/               - Информация за текущ потребител''')
    
    add_heading(doc, '3.3. Генериране на документи', 2)
    add_paragraph(doc, 'Процесът на генериране на актове включва:')
    add_bullet_point(doc, '1. Получаване на JSON контекст от frontend (име на проект, дати, имена и др.)')
    add_bullet_point(doc, '2. Зареждане на Word шаблон от media/templates/')
    add_bullet_point(doc, '3. Замяна на {{placeholder}} маркери с реални стойности')
    add_bullet_point(doc, '4. Премахване на незапълнени {{placeholder}} маркери')
    add_bullet_point(doc, '5. Запазване на генериран DOCX файл')
    add_bullet_point(doc, '6. Създаване на ZIP архив')
    add_bullet_point(doc, '7. Връщане на URL адреси към frontend')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Файл: backend/core/utils/document_generator.py', italic=True)
    
    doc.add_page_break()
    
    # 4. Frontend
    add_heading(doc, '4. Frontend (React + TypeScript)', 1)
    
    add_heading(doc, '4.1. Структура на приложението', 2)
    add_paragraph(doc, 'Frontend е изградено като Single Page Application (SPA) с React Router:')
    add_bullet_point(doc, '/ - Начална страница (Home)')
    add_bullet_point(doc, '/login - Страница за вход')
    add_bullet_point(doc, '/projects - Списък с проекти')
    add_bullet_point(doc, '/projects/new - Създаване на нов проект')
    add_bullet_point(doc, '/projects/:id - Детайли на проект')
    add_bullet_point(doc, '/documents - Генериране и управление на документи')
    add_bullet_point(doc, '/admin - Администраторски панел')
    
    add_heading(doc, '4.2. Ключови компоненти', 2)
    
    add_paragraph(doc, 'Header - Навигационна лента:', bold=True)
    add_bullet_point(doc, 'Лого и заглавие на системата')
    add_bullet_point(doc, 'Навигационно меню')
    add_bullet_point(doc, 'Потребителско меню (Профил/Изход)')
    add_bullet_point(doc, 'Мобилно-адаптивен дизайн')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Documents - Генериране на актове:', bold=True)
    add_bullet_point(doc, 'Форми за Акт 7, Акт 14, Акт 15')
    add_bullet_point(doc, 'Auto-fill от избран проект')
    add_bullet_point(doc, 'Date picker с dayjs')
    add_bullet_point(doc, 'Списък на създадени документи с timestamps')
    add_bullet_point(doc, 'Качване на документи')
    add_bullet_point(doc, 'Изтеглане на DOCX и ZIP файлове')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Projects - Управление на проекти:', bold=True)
    add_bullet_point(doc, 'Grid layout с карти')
    add_bullet_point(doc, 'Филтриране по статус')
    add_bullet_point(doc, 'Progress bar за завършеност')
    add_bullet_point(doc, 'Детайлна страница с пълна информация')
    
    add_heading(doc, '4.3. State Management', 2)
    add_paragraph(doc, 'Redux Toolkit се използва за глобално състояние:')
    add_bullet_point(doc, 'authSlice - Потребител, токен, admin права')
    add_bullet_point(doc, 'uiSlice - UI настройки, темна тема')
    
    add_heading(doc, '4.4. API комуникация', 2)
    add_paragraph(doc, 'Axios клиент (frontend/src/api/client.ts):')
    add_bullet_point(doc, 'Base URL: http://localhost:8000/api/')
    add_bullet_point(doc, 'withCredentials: true (за cookies)')
    add_bullet_point(doc, 'setAuthHeader() - Добавя Bearer token')
    add_bullet_point(doc, 'Автоматично error handling')
    
    add_heading(doc, '4.5. Мобилна адаптивност', 2)
    add_paragraph(doc, 'Имплементирана чрез:')
    add_bullet_point(doc, 'CSS media queries (@media max-width: 768px)')
    add_bullet_point(doc, 'Ant Design Grid система (xs/sm/md/lg breakpoints)')
    add_bullet_point(doc, 'Responsive колони: xs={24} md={8} (full width на mobile, 1/3 на desktop)')
    add_bullet_point(doc, 'Touch-friendly бутони (min-height: 44px)')
    add_bullet_point(doc, 'Оптимизирани padding и margins за mobile')
    
    doc.add_page_break()
    
    # 5. Генериране на документи
    add_heading(doc, '5. Генериране на документи - Детайлно', 1)
    
    add_heading(doc, '5.1. Word шаблони', 2)
    add_paragraph(doc, 'Системата използва официални български форми:')
    add_bullet_point(doc, 'act7_bg.docx - Акт за приемане на СМР по нива и елементи')
    add_bullet_point(doc, 'act14_bg.docx - Акт за приемане на конструкцията (Образец 14)')
    add_bullet_point(doc, 'act15_bg.docx - Констативен акт (Образец 15)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Всички шаблони са стандартизирани с:')
    add_bullet_point(doc, 'Шрифт: Times New Roman 12pt')
    add_bullet_point(doc, 'Placeholder синтаксис: {{field_name}}')
    add_bullet_point(doc, 'Премахнати текстови етикети преди placeholders')
    add_bullet_point(doc, 'Компактна структура без излишни празни редове')
    
    add_heading(doc, '5.2. Акт 7 - СМР по нива', 2)
    add_paragraph(doc, 'Полета за попълване:')
    add_bullet_point(doc, 'project_name - Име на строеж')
    add_bullet_point(doc, 'project_location - Местонахождение')
    add_bullet_point(doc, 'act_date - Дата на акт')
    add_bullet_point(doc, 'client_name - Възложител')
    add_bullet_point(doc, 'consultant_name - Консултант')
    add_bullet_point(doc, 'contractor_name - Строител')
    add_bullet_point(doc, 'designer_name - Проектант')
    add_bullet_point(doc, 'representative_builder - Представител на строителя')
    add_bullet_point(doc, 'representative_supervision - Представител на надзора')
    add_bullet_point(doc, 'representative_designer - Представител на проектанта')
    add_bullet_point(doc, 'level_from - Кота от')
    add_bullet_point(doc, 'level_to - Кота до')
    add_bullet_point(doc, 'work_description - Описание на работите')
    add_bullet_point(doc, 'execution - Забележки по изпълнение')
    
    add_heading(doc, '5.3. Акт 14 - Приемане на конструкцията', 2)
    add_paragraph(doc, 'Полета за попълване:')
    add_bullet_point(doc, 'act_date - Дата')
    add_bullet_point(doc, 'project_name - Строеж')
    add_bullet_point(doc, 'project_location - Местонахождение')
    add_bullet_point(doc, 'client_name - Възложител')
    add_bullet_point(doc, 'contractor_name - Строител')
    add_bullet_point(doc, 'designer_name - Проектант')
    add_bullet_point(doc, 'consultant_name - Консултант')
    add_bullet_point(doc, 'tech_supervisor_name - Технически ръководител')
    add_bullet_point(doc, 'additional_documents - Допълнителни документи')
    add_bullet_point(doc, 'defects_description - Описание на дефекти')
    
    add_heading(doc, '5.4. Акт 15 - Констативен акт', 2)
    add_paragraph(doc, 'Полета за попълване:')
    add_bullet_point(doc, 'act_date - Дата')
    add_bullet_point(doc, 'project_name - Строеж')
    add_bullet_point(doc, 'project_location - Местонахождение')
    add_bullet_point(doc, 'client_name - Възложител (в секция А)')
    add_bullet_point(doc, 'designer_name - Проектанти (в секция Б)')
    add_bullet_point(doc, 'contractor_name - Строител (в секция В)')
    add_bullet_point(doc, 'consultant_name - Консултант/Надзор (в секция Г)')
    add_bullet_point(doc, 'execution_findings - Констатации по изпълнението')
    add_bullet_point(doc, 'site_condition - Състояние на площадката')
    add_bullet_point(doc, 'surrounding_condition - Околно пространство')
    
    add_heading(doc, '5.5. Процес на генериране', 2)
    add_paragraph(doc, 'Стъпка по стъпка:')
    add_paragraph(doc, '')
    add_paragraph(doc, '1. Потребителят избира тип акт (7, 14 или 15)')
    add_paragraph(doc, '2. Опционално избира проект за auto-fill на полетата')
    add_paragraph(doc, '3. Попълва/редактира полетата във формата')
    add_paragraph(doc, '4. Натиска бутон "Генерирай [Акт X]"')
    add_paragraph(doc, '5. Frontend изпраща POST заявка към /api/documents/generate/')
    add_paragraph(doc, '6. Backend зарежда съответния шаблон')
    add_paragraph(doc, '7. Замяна на placeholders с данни от context')
    add_paragraph(doc, '8. Bidirectional mapping (consultant_name ↔ representative_supervision)')
    add_paragraph(doc, '9. Премахване на незапълнени {{placeholder}} с regex')
    add_paragraph(doc, '10. Запазване на DOCX файл в media/generated/')
    add_paragraph(doc, '11. Създаване на ZIP архив')
    add_paragraph(doc, '12. Връщане на JSON с URLs към frontend')
    add_paragraph(doc, '13. Показване в списъка на документи')
    
    doc.add_page_break()
    
    # 6. Модули
    add_heading(doc, '6. Модули и функционалности', 1)
    
    add_heading(doc, '6.1. Потребители и оторизация', 2)
    add_paragraph(doc, 'Системата използва Django authentication:')
    add_bullet_point(doc, 'Username/Password вход')
    add_bullet_point(doc, 'JWT Token authentication (Bearer token)')
    add_bullet_point(doc, 'Role-based access (Employee/Admin)')
    add_bullet_point(doc, 'Protected routes на frontend (ProtectedRoute компонент)')
    add_bullet_point(doc, 'Permission classes на backend (IsEmployeeOrAdmin)')
    
    add_heading(doc, '6.2. Управление на проекти', 2)
    add_bullet_point(doc, 'Създаване на нови строителни обекти')
    add_bullet_point(doc, 'Редактиране на информация')
    add_bullet_point(doc, 'Проследяване на напредък (progress bar 0-100%)')
    add_bullet_point(doc, 'Статус управление (active/completed/paused)')
    add_bullet_point(doc, 'Съхранение на ключова информация (възложител, проектант, строител, консултант)')
    
    add_heading(doc, '6.3. Документи и актове', 2)
    add_bullet_point(doc, 'Автоматично генериране на Word документи')
    add_bullet_point(doc, 'Съхранение с timestamps за идентификация')
    add_bullet_point(doc, 'Качване на външни документи')
    add_bullet_point(doc, 'Изтеглане на DOCX и ZIP файлове')
    add_bullet_point(doc, 'Изтриване на документи (с Popconfirm)')
    
    add_heading(doc, '6.4. Екипи и служители', 2)
    add_bullet_point(doc, 'Регистрация на служители')
    add_bullet_point(doc, 'Длъжности и права')
    add_bullet_point(doc, 'Администраторски достъп')
    add_bullet_point(doc, 'Контактна информация')
    
    doc.add_page_break()
    
    # 7. База данни
    add_heading(doc, '7. База данни', 1)
    
    add_heading(doc, '7.1. Schema', 2)
    add_paragraph(doc, 'Основни таблици:')
    add_bullet_point(doc, 'auth_user - Django потребители')
    add_bullet_point(doc, 'core_employee - Служители')
    add_bullet_point(doc, 'core_client - Клиенти/Възложители')
    add_bullet_point(doc, 'core_project - Проекти/Строежи')
    add_bullet_point(doc, 'core_document - Документи')
    add_bullet_point(doc, 'core_act - Актове')
    add_bullet_point(doc, 'core_task - Задачи')
    add_bullet_point(doc, 'core_projectmember - Участници в проект')
    
    add_heading(doc, '7.2. Връзки между таблици', 2)
    add_bullet_point(doc, 'Employee → User (OneToOne)')
    add_bullet_point(doc, 'Project → Client (ForeignKey)')
    add_bullet_point(doc, 'Act → Project (ForeignKey)')
    add_bullet_point(doc, 'ProjectMember → Project + Employee (ForeignKey)')
    add_bullet_point(doc, 'Task → Project (ForeignKey)')
    
    add_heading(doc, '7.3. Migrations', 2)
    add_paragraph(doc, 'Django migrations в backend/core/migrations/:')
    add_bullet_point(doc, '0001_initial.py - Първоначална схема')
    add_bullet_point(doc, '0002_*.py - Премахване на одобрения, добавяне на контекст')
    add_bullet_point(doc, '0003_add_project_fields.py - Нови полета за проекти')
    add_bullet_point(doc, '0004_*.py - Оптимизация на модели')
    add_bullet_point(doc, '0005_project_progress.py - Progress bar')
    add_bullet_point(doc, '0006_act.py - Модел Act')
    add_bullet_point(doc, '0007_zip_fields.py - ZIP файлове')
    
    doc.add_page_break()
    
    # 8. API документация
    add_heading(doc, '8. API документация', 1)
    
    add_heading(doc, '8.1. Authentication', 2)
    add_paragraph(doc, 'POST /api/auth/login/', bold=True)
    add_paragraph(doc, 'Request body:')
    add_code_block(doc, '''{
  "username": "admin",
  "password": "password123"
}''')
    add_paragraph(doc, 'Response:')
    add_code_block(doc, '''{
  "token": "eyJ0eXAiOiJKV1QiLCJhbGc...",
  "user": {
    "id": 1,
    "username": "admin",
    "is_admin": true
  }
}''')
    
    add_heading(doc, '8.2. Projects', 2)
    add_paragraph(doc, 'GET /api/projects/', bold=True)
    add_paragraph(doc, 'Response:')
    add_code_block(doc, '''[
  {
    "id": 1,
    "name": "Жилищна сграда - София",
    "location": "бул. Витоша 100",
    "status": "active",
    "progress": 65,
    "client": {...},
    "start_date": "2025-01-15",
    "end_date": "2025-12-31",
    "consultant_name": "Инж. Иван Петров",
    "contractor_name": "Строй ЕООД",
    "designer_name": "Архитект Мария Георгиева"
  }
]''')
    
    add_heading(doc, '8.3. Document Generation', 2)
    add_paragraph(doc, 'POST /api/documents/generate/', bold=True)
    add_paragraph(doc, 'Request body:')
    add_code_block(doc, '''{
  "template_name": "act14_bg.docx",
  "context": {
    "act_date": "28.11.2025",
    "project_name": "Жилищна сграда",
    "project_location": "гр. София",
    "client_name": "Строй Инвест ЕООД",
    "contractor_name": "Билд Груп АД",
    "designer_name": "Инж. Петър Димитров",
    "consultant_name": "Инж. Иван Петров"
  }
}''')
    add_paragraph(doc, 'Response:')
    add_code_block(doc, '''{
  "id": 15,
  "title": "act14_bg_20251128_143022.docx",
  "file_docx": "http://localhost:8000/media/generated/act14_bg_20251128_143022.docx",
  "zip_url": "http://localhost:8000/media/generated/act14_bg_20251128_143022.zip",
  "created_at": "2025-11-28T14:30:22.123Z"
}''')
    
    doc.add_page_break()
    
    # 9. Инсталация
    add_heading(doc, '9. Инсталация и стартиране', 1)
    
    add_heading(doc, '9.1. Предварителни изисквания', 2)
    add_bullet_point(doc, 'Python 3.8+ (препоръчва се 3.10+)')
    add_bullet_point(doc, 'Node.js 16+ (препоръчва се 18+)')
    add_bullet_point(doc, 'pip (Python package manager)')
    add_bullet_point(doc, 'npm (Node package manager)')
    add_bullet_point(doc, 'Git')
    
    add_heading(doc, '9.2. Backend setup', 2)
    add_paragraph(doc, 'Стъпки:')
    add_paragraph(doc, '')
    add_code_block(doc, '''# 1. Клониране на repo
git clone <repo-url>
cd construction-supervision

# 2. Създаване на виртуална среда
cd backend
python -m venv .venv
.venv\\Scripts\\activate  # Windows
source .venv/bin/activate  # Linux/Mac

# 3. Инсталиране на dependencies
pip install -r requirements.txt

# 4. Конфигурация (.env файл)
cp env.example .env
# Редактирайте .env с вашите настройки

# 5. Миграции на база данни
python manage.py migrate

# 6. Създаване на superuser
python manage.py createsuperuser

# 7. Стартиране на сървър
python manage.py runserver''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Backend ще бъде достъпен на: http://127.0.0.1:8000/')
    
    add_heading(doc, '9.3. Frontend setup', 2)
    add_code_block(doc, '''# 1. Отиваме във frontend папката
cd frontend

# 2. Инсталиране на dependencies
npm install

# 3. Стартиране на dev server
npm run dev''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Frontend ще бъде достъпен на: http://localhost:5173/ или http://localhost:5174/')
    
    add_heading(doc, '9.4. Бърз старт със scripts', 2)
    add_paragraph(doc, 'За Windows PowerShell:')
    add_code_block(doc, '''powershell -ExecutionPolicy Bypass -File scripts\\start-dev.ps1''')
    add_paragraph(doc, 'Този скрипт автоматично стартира и backend, и frontend.')
    
    doc.add_page_break()
    
    # 10. Deployment
    add_heading(doc, '10. Deployment (Production)', 1)
    
    add_heading(doc, '10.1. Backend deployment', 2)
    add_paragraph(doc, 'Препоръчителни платформи:')
    add_bullet_point(doc, 'Railway.app - Безплатен tier, автоматичен deployment от GitHub')
    add_bullet_point(doc, 'Render.com - Безплатен tier, PostgreSQL включен')
    add_bullet_point(doc, 'AWS EC2 - Пълен контрол, по-скъпо')
    add_bullet_point(doc, 'DigitalOcean - Droplets, managed databases')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Production настройки (settings.py):')
    add_bullet_point(doc, 'DEBUG = False')
    add_bullet_point(doc, 'ALLOWED_HOSTS = ["yourdomain.com"]')
    add_bullet_point(doc, 'PostgreSQL database')
    add_bullet_point(doc, 'Gunicorn WSGI server')
    add_bullet_point(doc, 'Static files на CDN')
    add_bullet_point(doc, 'HTTPS/SSL сертификат')
    
    add_heading(doc, '10.2. Frontend deployment', 2)
    add_paragraph(doc, 'Препоръчителни платформи:')
    add_bullet_point(doc, 'Vercel - Най-добро за React/Vite, безплатно')
    add_bullet_point(doc, 'Netlify - Безплатно, CI/CD включен')
    add_bullet_point(doc, 'Cloudflare Pages - Бързо CDN')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Build команди:')
    add_code_block(doc, '''npm run build
# Генерира production build в dist/''')
    
    add_heading(doc, '10.3. Environment variables', 2)
    add_paragraph(doc, 'Backend .env:')
    add_code_block(doc, '''DEBUG=False
SECRET_KEY=your-secret-key-here
DATABASE_URL=postgresql://user:pass@host:5432/dbname
ALLOWED_HOSTS=yourdomain.com
CORS_ALLOWED_ORIGINS=https://yourfrontend.com''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Frontend environment:')
    add_code_block(doc, '''VITE_API_URL=https://yourbackend.com/api/''')
    
    doc.add_page_break()
    
    # 11. Примери
    add_heading(doc, '11. Примери за използване', 1)
    
    add_heading(doc, '11.1. Създаване на нов проект', 2)
    add_paragraph(doc, '1. Влезте в системата с username/password')
    add_paragraph(doc, '2. Отидете на страница "Обекти"')
    add_paragraph(doc, '3. Натиснете бутон "Нов обект"')
    add_paragraph(doc, '4. Попълнете форма:')
    add_bullet_point(doc, '  • Име на проект: "Жилищна сграда - ул. Оборище 15"')
    add_bullet_point(doc, '  • Местонахождение: "гр. София, ул. Оборище 15"')
    add_bullet_point(doc, '  • Възложител: Изберете от списък или създайте нов')
    add_bullet_point(doc, '  • Начална дата: 01.12.2025')
    add_bullet_point(doc, '  • Крайна дата: 30.06.2026')
    add_bullet_point(doc, '  • Консултант: "Инж. Петър Георгиев"')
    add_bullet_point(doc, '  • Проектант: "Архитект Мария Иванова"')
    add_bullet_point(doc, '  • Строител: "Строй Компани ЕООД"')
    add_paragraph(doc, '5. Натиснете "Създай"')
    add_paragraph(doc, '6. Проектът се появява в списъка')
    
    add_heading(doc, '11.2. Генериране на Акт 14', 2)
    add_paragraph(doc, '1. Отидете на страница "Документи"')
    add_paragraph(doc, '2. В секцията "Изберете обект за авто-попълване" изберете проект')
    add_paragraph(doc, '3. Полетата автоматично се попълват')
    add_paragraph(doc, '4. Проверете/редактирайте полетата:')
    add_bullet_point(doc, '  • Дата на акт: 28.11.2025')
    add_bullet_point(doc, '  • Технически ръководител: "Инж. Стоян Димитров"')
    add_bullet_point(doc, '  • Допълнителни документи: "Протокол №5, Декларация за съответствие"')
    add_bullet_point(doc, '  • Дефекти: "Няма установени дефекти"')
    add_paragraph(doc, '5. Натиснете "Генерирай Акт 14"')
    add_paragraph(doc, '6. Документът се появява в списъка долу')
    add_paragraph(doc, '7. Изтеглете DOCX или ZIP файл')
    
    add_heading(doc, '11.3. Качване на собствен документ', 2)
    add_paragraph(doc, '1. На страница "Документи"')
    add_paragraph(doc, '2. Натиснете бутон "Качи документ" (в секция "Списък на документи")')
    add_paragraph(doc, '3. Попълнете:')
    add_bullet_point(doc, '  • Заглавие: "Снимки от обект - 28.11.2025"')
    add_bullet_point(doc, '  • DOCX файл: Изберете файл от компютъра')
    add_paragraph(doc, '4. Натиснете "Качи"')
    add_paragraph(doc, '5. Файлът се появява в списъка')
    
    add_heading(doc, '11.4. Проследяване на напредък', 2)
    add_paragraph(doc, '1. Отидете на детайлна страница на проект')
    add_paragraph(doc, '2. Кликнете "Редактирай"')
    add_paragraph(doc, '3. Променете "Напредък на проекта" на 75%')
    add_paragraph(doc, '4. Запазете промените')
    add_paragraph(doc, '5. Progress bar се обновява визуално')
    
    doc.add_page_break()
    
    # Заключение
    add_heading(doc, 'Заключение', 1)
    add_paragraph(doc, 'Системата за Строителен Надзор е модерна, скалируема платформа, която автоматизира създаването на официални строителни документи и улеснява управлението на проекти.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Ключови предимства:', bold=True)
    add_bullet_point(doc, 'Автоматично генериране на актове с официални български форми')
    add_bullet_point(doc, 'Централизирано управление на всички строителни обекти')
    add_bullet_point(doc, 'Мобилно-адаптивен интерфейс за работа от всяко място')
    add_bullet_point(doc, 'Сигурна автентикация и контрол на достъпа')
    add_bullet_point(doc, 'Лесно deployment в cloud платформи')
    add_bullet_point(doc, 'Отворен код и възможност за разширяване')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Бъдещи подобрения:', bold=True)
    add_bullet_point(doc, 'PDF генериране директно от backend')
    add_bullet_point(doc, 'Електронни подписи на документи')
    add_bullet_point(doc, 'Мобилно приложение (iOS/Android)')
    add_bullet_point(doc, 'Напреднали отчети и анализи')
    add_bullet_point(doc, 'Интеграция с външни системи')
    add_bullet_point(doc, 'Real-time уведомления')
    
    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'За въпроси и поддръжка: support@construction-supervision.bg', italic=True)
    add_paragraph(doc, 'GitHub: https://github.com/your-username/construction-supervision', italic=True)
    
    # Добавяне на секция за скрийншоти
    doc.add_page_break()
    add_heading(doc, 'Приложение: Скрийншоти', 1)
    add_paragraph(doc, 'Моля, добавете следните скрийншоти към този документ:', italic=True)
    add_paragraph(doc, '')
    add_paragraph(doc, '1. Начална страница (Home)')
    add_paragraph(doc, '   - Показва header, hero секция, информация за екипа')
    add_paragraph(doc, '')
    add_paragraph(doc, '2. Страница за вход (Login)')
    add_paragraph(doc, '   - Форма за username/password')
    add_paragraph(doc, '')
    add_paragraph(doc, '3. Списък с проекти (Projects)')
    add_paragraph(doc, '   - Grid layout с карти, progress bars')
    add_paragraph(doc, '')
    add_paragraph(doc, '4. Детайли на проект')
    add_paragraph(doc, '   - Пълна информация, бутони за редакция')
    add_paragraph(doc, '')
    add_paragraph(doc, '5. Страница Документи - Форми за актове')
    add_paragraph(doc, '   - Три колони с Акт 7, 14, 15')
    add_paragraph(doc, '')
    add_paragraph(doc, '6. Страница Документи - Списък на документи')
    add_paragraph(doc, '   - Списък с timestamps, download links')
    add_paragraph(doc, '')
    add_paragraph(doc, '7. Генериран Акт 14 в Word')
    add_paragraph(doc, '   - Отворен DOCX файл с попълнени данни')
    add_paragraph(doc, '')
    add_paragraph(doc, '8. Мобилна версия')
    add_paragraph(doc, '   - Header на mobile, форма на малък екран')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Забележка: Можете да вмъкнете скрийншоти с Insert > Pictures в Microsoft Word')
    
    # Запазване на документа
    output_path = 'Документация_Система_за_Строителен_Надзор.docx'
    doc.save(output_path)
    print(f'✓ Документацията е създадена: {output_path}')
    print(f'✓ Файлът съдържа {len(doc.paragraphs)} параграфа')
    print(f'✓ Моля, отворете файла и добавете скрийншоти в секция "Приложение: Скрийншоти"')
    
    return output_path

if __name__ == '__main__':
    create_documentation()
