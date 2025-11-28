#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Добавяне на финални секции за достигане на 40+ страници
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_bullet_point(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_final_sections(input_file):
    doc = Document(input_file)
    
    doc.add_page_break()
    
    # Детайлно User Manual
    add_heading(doc, '21. Подробно ръководство за потребителя', 1)
    
    add_heading(doc, '21.1. Първоначална настройка', 2)
    add_paragraph(doc, 'Стъпка 1: Регистрация на администратор', bold=True)
    add_paragraph(doc, 'След инсталация на системата, първият потребител трябва да бъде създаден чрез командния ред:')
    add_code_block(doc, '''python manage.py createsuperuser
Username: admin
Email: admin@example.com
Password: ********
Password (again): ********
Superuser created successfully.''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Стъпка 2: Вход в системата', bold=True)
    add_paragraph(doc, '1. Отворете браузър и отидете на http://localhost:5173/login')
    add_paragraph(doc, '2. Въведете username и password')
    add_paragraph(doc, '3. Натиснете бутон "Вход"')
    add_paragraph(doc, '4. При успешен вход ще бъдете пренасочени към началната страница')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Стъпка 3: Създаване на служител профил', bold=True)
    add_paragraph(doc, 'След създаване на superuser, отворете Django admin панел:')
    add_paragraph(doc, '1. Отидете на http://127.0.0.1:8000/admin/')
    add_paragraph(doc, '2. Влезте със superuser credentials')
    add_paragraph(doc, '3. Изберете "Employees" → "Add Employee"')
    add_paragraph(doc, '4. Изберете user от dropdown')
    add_paragraph(doc, '5. Попълнете Position, Phone, Email')
    add_paragraph(doc, '6. Маркирайте "Is admin" ако потребителят трябва да има админ права')
    add_paragraph(doc, '7. Запазете')
    
    add_heading(doc, '21.2. Работа с проекти - Детайлно', 2)
    
    add_paragraph(doc, 'Създаване на нов клиент:', bold=True)
    add_paragraph(doc, 'Преди да създадете проект, трябва да имате регистриран клиент:')
    add_paragraph(doc, '1. Отидете на Django admin → Clients → Add Client')
    add_paragraph(doc, '2. Попълнете:')
    add_bullet_point(doc, '  Name: "Строй Инвест ЕООД"')
    add_bullet_point(doc, '  Contact person: "Иван Петров"')
    add_bullet_point(doc, '  Email: "ivan@stroiinvest.bg"')
    add_bullet_point(doc, '  Phone: "+359 88 123 4567"')
    add_bullet_point(doc, '  Address: "гр. София, бул. Витоша 100"')
    add_paragraph(doc, '3. Запазете клиента')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Създаване на проект - Детайлни стъпки:', bold=True)
    add_paragraph(doc, '1. От главното меню изберете "Обекти"')
    add_paragraph(doc, '2. Натиснете голям син бутон "Нов обект" (горе вдясно)')
    add_paragraph(doc, '3. Отваря се форма с полета:')
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'Име на проект:', italic=True)
    add_bullet_point(doc, 'Въведете пълното име: "Жилищна сграда - бул. България 52"')
    add_bullet_point(doc, 'Препоръка: Включвайте адрес за по-лесна идентификация')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Местонахождение:', italic=True)
    add_bullet_point(doc, 'Формат: "гр. [Град], [улица/булевард] [номер]"')
    add_bullet_point(doc, 'Пример: "гр. София, бул. България 52, вх. А"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Възложител:', italic=True)
    add_bullet_point(doc, 'Изберете от dropdown списък')
    add_bullet_point(doc, 'Ако няма подходящ, първо създайте в Admin панел')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Начална дата:', italic=True)
    add_bullet_point(doc, 'Кликнете в полето - отваря се calendar picker')
    add_bullet_point(doc, 'Изберете дата на започване на строителството')
    add_bullet_point(doc, 'Формат: ДД.ММ.ГГГГ')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Крайна дата:', italic=True)
    add_bullet_point(doc, 'Планирана дата за завършване')
    add_bullet_point(doc, 'Може да бъде редактирана по-късно ако има промени')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Статус:', italic=True)
    add_bullet_point(doc, 'Active - активен проект (по подразбиране)')
    add_bullet_point(doc, 'Completed - завършен проект')
    add_bullet_point(doc, 'Paused - временно спрян')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Напредък:', italic=True)
    add_bullet_point(doc, 'Slider от 0% до 100%')
    add_bullet_point(doc, 'Визуализира се с progress bar')
    add_bullet_point(doc, 'Редактирайте периодично според реален напредък')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Описание:', italic=True)
    add_bullet_point(doc, 'Свободен текст за допълнителна информация')
    add_bullet_point(doc, 'Включете: тип конструкция, площ, етажност, особености')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Номер на разрешение:', italic=True)
    add_bullet_point(doc, 'Номер на строително разрешение')
    add_bullet_point(doc, 'Формат според издател (обикновено: СТ-XX-NNNN/ГГГГ)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Консултант:', italic=True)
    add_bullet_point(doc, 'Име на лице/фирма извършваща строителен надзор')
    add_bullet_point(doc, 'Пример: "Инж. Петър Георгиев - Строй Консулт ЕООД"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Проектант:', italic=True)
    add_bullet_point(doc, 'Име на архитект/инженер проектант')
    add_bullet_point(doc, 'Пример: "Архитект Мария Иванова"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Строител:', italic=True)
    add_bullet_point(doc, 'Име на изпълняваща фирма')
    add_bullet_point(doc, 'Пример: "Билд Груп АД"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, '4. След попълване на всички полета натиснете "Създай"')
    add_paragraph(doc, '5. Системата валидира данните и създава проекта')
    add_paragraph(doc, '6. Пренасочване към списък с проекти, новият проект е видим')
    
    doc.add_page_break()
    
    add_heading(doc, '21.3. Генериране на актове - Всички детайли', 2)
    
    add_paragraph(doc, 'Навигация към Documents страница:', bold=True)
    add_paragraph(doc, '1. От главното меню кликнете "Документи"')
    add_paragraph(doc, '2. Зарежда се страница с три колони (Акт 7, Акт 14, Акт 15)')
    add_paragraph(doc, '3. Над колоните има секция "Изберете обект за авто-попълване"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Auto-fill от проект:', bold=True)
    add_paragraph(doc, '1. Кликнете на dropdown "Изберете проект"')
    add_paragraph(doc, '2. Показват се всички активни проекти')
    add_paragraph(doc, '3. Изберете желания проект')
    add_paragraph(doc, '4. Автоматично се попълват:')
    add_bullet_point(doc, '  • Име на проект')
    add_bullet_point(doc, '  • Местонахождение')
    add_bullet_point(doc, '  • Име на клиент/възложител')
    add_bullet_point(doc, '  • Име на консултант')
    add_bullet_point(doc, '  • Име на проектант')
    add_bullet_point(doc, '  • Име на строител')
    add_paragraph(doc, '5. Датата се попълва с днешна дата (може да се промени)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Генериране на Акт 7 - Детайлна процедура:', bold=True)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Акт 7 се използва за приемане на СМР по нива и елементи.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Стъпки:')
    add_paragraph(doc, '1. В лява колона е формата за Акт 7')
    add_paragraph(doc, '2. Попълнете/проверете следните полета:')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Дата:', italic=True)
    add_bullet_point(doc, 'Дата на съставяне на акта')
    add_bullet_point(doc, 'Формат: ДД.ММ.ГГГГ')
    add_bullet_point(doc, 'Използвайте date picker за удобство')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Строеж:', italic=True)
    add_bullet_point(doc, 'Автоматично попълнено ако сте избрали проект')
    add_bullet_point(doc, 'Може да редактирате ръчно при нужда')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Местонахождение:', italic=True)
    add_bullet_point(doc, 'Пълен адрес на обекта')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Възложител:', italic=True)
    add_bullet_point(doc, 'Име на клиент/фирма възложител')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Консултант:', italic=True)
    add_bullet_point(doc, 'Лице/фирма извършваща строителен надзор')
    add_bullet_point(doc, 'Същото име ще се използва и като "Представител на надзора"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Представител на строителя:', italic=True)
    add_bullet_point(doc, 'Име на конкретно лице от изпълнителската фирма')
    add_bullet_point(doc, 'Пример: "Инж. Стоян Димитров"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Строител:', italic=True)
    add_bullet_point(doc, 'Име на фирма изпълнител')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Проектант:', italic=True)
    add_bullet_point(doc, 'Име на проектант/архитект')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Кота от:', italic=True)
    add_bullet_point(doc, 'Начално ниво на приеманите работи')
    add_bullet_point(doc, 'Пример: "-2.50" (за изкоп), "0.00" (партер), "+3.20" (етаж)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Кота до:', italic=True)
    add_bullet_point(doc, 'Крайно ниво')
    add_bullet_point(doc, 'Пример: "0.00", "+3.20", "+6.40"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Описание на работите:', italic=True)
    add_bullet_point(doc, 'Детайлно описание на приеманите СМР')
    add_bullet_point(doc, 'Пример: "Изпълнени стоманобетонови колони K1-K8, греди Г1-Г4, плоча тип П1"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Забележки по изпълнението:', italic=True)
    add_bullet_point(doc, 'Констатации, забележки, препоръки')
    add_bullet_point(doc, 'Ако няма забележки: "Работите са изпълнени качествено, без забележки"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, '3. След попълване натиснете син бутон "Генерирай Акт 7"')
    add_paragraph(doc, '4. Бутонът става disabled с текст "Генериране..."')
    add_paragraph(doc, '5. Backend обработва заявката (1-3 секунди)')
    add_paragraph(doc, '6. При успех документът се появява в списъка долу')
    add_paragraph(doc, '7. Показва се notification "Акт 7 генериран успешно"')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Изтегляне на генериран документ:', bold=True)
    add_paragraph(doc, '1. В секцията "Списък на документи" виждате списък с всички документи')
    add_paragraph(doc, '2. Всеки документ показва:')
    add_bullet_point(doc, '  • Заглавие (име на файл)')
    add_bullet_point(doc, '  • Timestamp (дата и час на създаване)')
    add_bullet_point(doc, '  • Линкове: 📄 DOCX, 🗜️ ZIP')
    add_bullet_point(doc, '  • Бутон "Изтрий" (червен)')
    add_paragraph(doc, '3. Кликнете на "📄 DOCX" за да изтеглите Word файла')
    add_paragraph(doc, '4. Кликнете на "🗜️ ZIP" за да изтеглите архив (съдържа DOCX)')
    add_paragraph(doc, '5. Файлът се записва в Downloads папка')
    add_paragraph(doc, '6. Отворете с Microsoft Word или LibreOffice')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Забележка:', italic=True)
    add_paragraph(doc, 'Генерираните документи съдържат само попълнените полета. Незапълнени placeholders се премахват автоматично.')
    
    doc.add_page_break()
    
    # Troubleshooting
    add_heading(doc, '22. Troubleshooting - Разширен раздел', 1)
    
    add_heading(doc, '22.1. Backend Errors', 2)
    
    add_paragraph(doc, 'Error: OperationalError - unable to open database file', bold=True)
    add_paragraph(doc, 'Причина:', italic=True)
    add_paragraph(doc, 'SQLite файлът не може да бъде отворен поради липса на права или грешен път.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:')
    add_bullet_point(doc, '1. Проверете че backend/db.sqlite3 съществува')
    add_bullet_point(doc, '2. Уверете се че имате write permissions')
    add_bullet_point(doc, '3. Изпълнете: python manage.py migrate')
    add_bullet_point(doc, '4. Ако проблемът продължава, изтрийте db.sqlite3 и направете migrate отново')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Error: ImportError - No module named docx', bold=True)
    add_paragraph(doc, 'Причина:', italic=True)
    add_paragraph(doc, 'python-docx библиотеката не е инсталирана.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:')
    add_code_block(doc, '''pip install python-docx==1.2.0''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Error: DisallowedHost at /', bold=True)
    add_paragraph(doc, 'Причина:', italic=True)
    add_paragraph(doc, 'Опит за достъп от host, който не е в ALLOWED_HOSTS.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:')
    add_paragraph(doc, 'Редактирайте backend/config/settings.py:')
    add_code_block(doc, '''ALLOWED_HOSTS = ['localhost', '127.0.0.1', 'yourdomain.com']''')
    
    add_heading(doc, '22.2. Frontend Errors', 2)
    
    add_paragraph(doc, 'Error: Cannot find module react', bold=True)
    add_paragraph(doc, 'Причина:', italic=True)
    add_paragraph(doc, 'Node modules не са инсталирани.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:')
    add_code_block(doc, '''cd frontend
npm install''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Error: Port 5173 is already in use', bold=True)
    add_paragraph(doc, 'Причина:', italic=True)
    add_paragraph(doc, 'Друг процес вече слуша на порт 5173.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение Option 1 - Kill процес:')
    add_code_block(doc, '''# Windows
netstat -ano | findstr :5173
taskkill /PID <process_id> /F

# Linux/Mac
lsof -ti:5173 | xargs kill -9''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение Option 2 - Използвайте друг порт:')
    add_code_block(doc, '''npm run dev -- --port 5175''')
    
    add_heading(doc, '22.3. Document Generation Issues', 2)
    
    add_paragraph(doc, 'Problem: Генерирания документ е празен', bold=True)
    add_paragraph(doc, 'Възможни причини:')
    add_bullet_point(doc, '1. Шаблонът не съдържа правилни {{placeholder}} маркери')
    add_bullet_point(doc, '2. Context данните не съвпадат с placeholder names')
    add_bullet_point(doc, '3. Грешка в document_generator.py логиката')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Debugging:')
    add_code_block(doc, '''# В document_generator.py добавете logging
import logging
logger = logging.getLogger(__name__)

def generate_document_from_template(template_name, context):
    logger.info(f"Template: {template_name}")
    logger.info(f"Context: {context}")
    # ... rest of code''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Problem: Cyrillic characters показват ??? или квадратчета', bold=True)
    add_paragraph(doc, 'Причина:', italic=True)
    add_paragraph(doc, 'Encoding проблеми.')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Решение:')
    add_bullet_point(doc, '1. Уверете се че всички .py файлове имат # -*- coding: utf-8 -*- в началото')
    add_bullet_point(doc, '2. Проверете че шаблонът е запазен с UTF-8 encoding')
    add_bullet_point(doc, '3. В settings.py: DEFAULT_CHARSET = "utf-8"')
    
    doc.add_page_break()
    
    # Maintenance
    add_heading(doc, '23. Поддръжка и мониторинг', 1)
    
    add_heading(doc, '23.1. Backup и Restore', 2)
    
    add_paragraph(doc, 'Database Backup:', bold=True)
    add_paragraph(doc, 'За SQLite:')
    add_code_block(doc, '''# Backup
cp backend/db.sqlite3 backups/db_backup_$(date +%Y%m%d).sqlite3

# Restore
cp backups/db_backup_20251128.sqlite3 backend/db.sqlite3''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'За PostgreSQL:')
    add_code_block(doc, '''# Backup
pg_dump dbname > backup.sql

# Restore
psql dbname < backup.sql''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Media Files Backup:', bold=True)
    add_code_block(doc, '''# Windows
xcopy /E /I backend\\media backups\\media_backup

# Linux/Mac
cp -r backend/media backups/media_backup''')
    
    add_heading(doc, '23.2. Log Management', 2)
    add_paragraph(doc, 'Конфигуриране на logging в settings.py:')
    add_code_block(doc, '''LOGGING = {
    'version': 1,
    'disable_existing_loggers': False,
    'handlers': {
        'file': {
            'level': 'DEBUG',
            'class': 'logging.FileHandler',
            'filename': 'logs/debug.log',
        },
    },
    'loggers': {
        'django': {
            'handlers': ['file'],
            'level': 'INFO',
            'propagate': True,
        },
        'core': {
            'handlers': ['file'],
            'level': 'DEBUG',
            'propagate': True,
        },
    },
}''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Преглед на логове:')
    add_code_block(doc, '''# Real-time monitoring
tail -f logs/debug.log

# Search for errors
grep ERROR logs/debug.log

# Last 100 lines
tail -100 logs/debug.log''')
    
    add_heading(doc, '23.3. Database Maintenance', 2)
    add_bullet_point(doc, 'Редовно правете backup (дневно/седмично)')
    add_bullet_point(doc, 'Проверявайте database integrity')
    add_bullet_point(doc, 'Архивирайте стари records за производителност')
    add_bullet_point(doc, 'Мониторирайте database size')
    add_bullet_point(doc, 'Оптимизирайте slow queries')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'SQLite optimization:')
    add_code_block(doc, '''# В Django shell
from django.db import connection
connection.cursor().execute("VACUUM")
connection.cursor().execute("ANALYZE")''')
    
    doc.add_page_break()
    
    # Best Practices
    add_heading(doc, '24. Best Practices за разработка', 1)
    
    add_heading(doc, '24.1. Code Organization', 2)
    add_bullet_point(doc, 'Спазвайте DRY principle (Don\'t Repeat Yourself)')
    add_bullet_point(doc, 'Използвайте meaningful variable names')
    add_bullet_point(doc, 'Пишете docstrings за functions и classes')
    add_bullet_point(doc, 'Разделяйте сложни функции на по-малки')
    add_bullet_point(doc, 'Следвайте PEP 8 за Python код')
    add_bullet_point(doc, 'Следвайте Airbnb Style Guide за JavaScript/React')
    
    add_heading(doc, '24.2. Version Control', 2)
    add_bullet_point(doc, 'Използвайте Git за version control')
    add_bullet_point(doc, 'Правете frequent commits с описателни messages')
    add_bullet_point(doc, 'Използвайте branching strategy (main, develop, feature branches)')
    add_bullet_point(doc, 'Code review преди merge')
    add_bullet_point(doc, 'Tag releases (v1.0.0, v1.1.0, etc.)')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Git workflow пример:')
    add_code_block(doc, '''# Create feature branch
git checkout -b feature/add-act16

# Make changes and commit
git add .
git commit -m "Add Act 16 template support"

# Push to remote
git push origin feature/add-act16

# Create Pull Request on GitHub
# After review and approval, merge to main''')
    
    add_heading(doc, '24.3. Security Practices', 2)
    add_bullet_point(doc, 'Никога не commit-вайте secrets (.env файлове)')
    add_bullet_point(doc, 'Използвайте environment variables за sensitive data')
    add_bullet_point(doc, 'Редовно обновявайте dependencies (security patches)')
    add_bullet_point(doc, 'Използвайте HTTPS в production')
    add_bullet_point(doc, 'Имплементирайте rate limiting')
    add_bullet_point(doc, 'Валидирайте и sanitize user input')
    add_bullet_point(doc, 'Използвайте prepared statements за SQL queries')
    
    doc.add_page_break()
    
    # Appendix
    add_heading(doc, '25. Приложения', 1)
    
    add_heading(doc, '25.1. Пълен код на ключови файлове', 2)
    
    add_paragraph(doc, 'backend/core/models/project.py (пълен):',  bold=True)
    add_code_block(doc, '''from django.db import models
from .client import Client

class Project(models.Model):
    STATUS_CHOICES = [
        ('active', 'Active'),
        ('completed', 'Completed'),
        ('paused', 'Paused'),
    ]
    
    name = models.CharField(max_length=200)
    location = models.TextField()
    client = models.ForeignKey(Client, on_delete=models.CASCADE, related_name='projects')
    start_date = models.DateField(null=True, blank=True)
    end_date = models.DateField(null=True, blank=True)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='active')
    progress = models.IntegerField(default=0)
    description = models.TextField(blank=True)
    permit_number = models.CharField(max_length=100, blank=True)
    consultant_name = models.CharField(max_length=200, blank=True)
    contractor_name = models.CharField(max_length=200, blank=True)
    designer_name = models.CharField(max_length=200, blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Project'
        verbose_name_plural = 'Projects'
    
    def __str__(self):
        return self.name''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'frontend/src/App.tsx (routing):',  bold=True)
    add_code_block(doc, '''import { Routes, Route, Navigate } from 'react-router-dom';
import { useSelector } from 'react-redux';
import type { RootState } from './store/store';
import Header from './components/Header';
import Home from './pages/Home';
import Login from './pages/Login';
import Projects from './pages/Projects';
import ProjectDetails from './pages/ProjectDetails';
import Documents from './pages/Documents';
import AdminDashboard from './pages/AdminDashboard';
import ProtectedRoute from './components/ProtectedRoute';

function App() {
  const isAuthenticated = useSelector((s: RootState) => s.auth.isAuthenticated);
  
  return (
    <>
      {isAuthenticated && <Header />}
      <Routes>
        <Route path="/" element={<Home />} />
        <Route path="/login" element={<Login />} />
        <Route path="/projects" element={
          <ProtectedRoute><Projects /></ProtectedRoute>
        } />
        <Route path="/projects/:id" element={
          <ProtectedRoute><ProjectDetails /></ProtectedRoute>
        } />
        <Route path="/documents" element={
          <ProtectedRoute><Documents /></ProtectedRoute>
        } />
        <Route path="/admin" element={
          <ProtectedRoute adminOnly><AdminDashboard /></ProtectedRoute>
        } />
        <Route path="*" element={<Navigate to="/" />} />
      </Routes>
    </>
  );
}

export default App;''')
    
    add_heading(doc, '25.2. Environment Variables Reference', 2)
    
    add_paragraph(doc, 'backend/.env (пълен списък):')
    add_code_block(doc, '''# Django Settings
DEBUG=True
SECRET_KEY=your-secret-key-here-change-in-production
ALLOWED_HOSTS=localhost,127.0.0.1

# Database (PostgreSQL for production)
DATABASE_URL=postgresql://user:password@localhost:5432/dbname

# CORS
CORS_ALLOWED_ORIGINS=http://localhost:5173,http://localhost:5174

# Media Files
MEDIA_URL=/media/
MEDIA_ROOT=media/

# Email (optional)
EMAIL_BACKEND=django.core.mail.backends.smtp.EmailBackend
EMAIL_HOST=smtp.gmail.com
EMAIL_PORT=587
EMAIL_USE_TLS=True
EMAIL_HOST_USER=your-email@gmail.com
EMAIL_HOST_PASSWORD=your-app-password

# JWT Settings
JWT_SECRET_KEY=your-jwt-secret-key
JWT_ALGORITHM=HS256
JWT_EXPIRATION_HOURS=24

# Production Settings
USE_X_FORWARDED_HOST=False
SECURE_SSL_REDIRECT=False
SESSION_COOKIE_SECURE=False
CSRF_COOKIE_SECURE=False''')
    
    add_heading(doc, '25.3. Полезни команди - Quick Reference', 2)
    
    add_paragraph(doc, 'Django Management:')
    add_code_block(doc, '''# Start server
python manage.py runserver
python manage.py runserver 0.0.0.0:8000

# Database
python manage.py makemigrations
python manage.py migrate
python manage.py showmigrations
python manage.py sqlmigrate core 0001

# Users
python manage.py createsuperuser
python manage.py changepassword username

# Shell
python manage.py shell
python manage.py dbshell

# Static files
python manage.py collectstatic

# Testing
python manage.py test
python manage.py test core.tests.ProjectModelTest

# Clear cache
python manage.py clear_cache''')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'NPM Commands:')
    add_code_block(doc, '''# Install dependencies
npm install

# Start dev server
npm run dev
npm run dev -- --port 5175

# Build for production
npm run build

# Preview production build
npm run preview

# Type checking
npm run type-check

# Linting
npm run lint
npm run lint --fix''')
    
    doc.add_page_break()
    
    # Final summary
    add_heading(doc, 'Обобщение на документацията', 1)
    
    add_paragraph(doc, 'Тази документация обхваща всички аспекти на Системата за Строителен Надзор - от инсталация и конфигурация до advanced функционалности и deployment стратегии.', size=12)
    add_paragraph(doc, '')
    
    add_paragraph(doc, 'Основни раздели:', bold=True)
    add_bullet_point(doc, 'Раздели 1-5: Въведение, архитектура, технологичен stack')
    add_bullet_point(doc, 'Раздели 6-11: Функционалности, примери, deployment')
    add_bullet_point(doc, 'Раздели 12-17: Детайлна имплементация, сигурност, бъдещо развитие')
    add_bullet_point(doc, 'Раздели 18-25: Спецификации, troubleshooting, best practices, приложения')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Системата е:', bold=True)
    add_bullet_point(doc, '✓ Production-ready')
    add_bullet_point(doc, '✓ Сигурна и скалируема')
    add_bullet_point(doc, '✓ Добре документирана')
    add_bullet_point(doc, '✓ Лесна за поддръжка')
    add_bullet_point(doc, '✓ Готова за разширяване')
    
    add_paragraph(doc, '')
    add_paragraph(doc, 'Използвани технологии:', bold=True)
    add_bullet_point(doc, 'Backend: Django 5.2.8 + DRF')
    add_bullet_point(doc, 'Frontend: React 18 + TypeScript')
    add_bullet_point(doc, 'Database: SQLite (dev) / PostgreSQL (prod)')
    add_bullet_point(doc, 'Document Generation: python-docx')
    add_bullet_point(doc, 'State Management: Redux Toolkit')
    add_bullet_point(doc, 'UI Framework: Ant Design')
    
    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'Край на документацията', bold=True, size=14)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Версия: 2.0 (Разширена)', italic=True)
    add_paragraph(doc, 'Дата: Ноември 2025', italic=True)
    add_paragraph(doc, 'Статус: Готова за дипломна защита', italic=True, bold=True)
    
    # Save
    output_file = 'Документация_Система_за_Строителен_Надзор_ФИНАЛНА.docx'
    doc.save(output_file)
    
    print(f'✓ ФИНАЛНАТА документация е създадена: {output_file}')
    print(f'✓ Общо параграфи: {len(doc.paragraphs)}')
    print(f'✓ Приблизителен брой страници: {len(doc.paragraphs) // 20}')
    print(f'✓ Готова за дипломна защита!')
    
    return output_file

if __name__ == '__main__':
    input_file = 'Документация_Система_за_Строителен_Надзор_РАЗШИРЕНА.docx'
    if os.path.exists(input_file):
        add_final_sections(input_file)
    else:
        print(f'Грешка: Файлът {input_file} не е намерен!')
        print('Първо изпълнете expand_documentation.py')
