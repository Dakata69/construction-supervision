#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convert PRESENTATION_EXPLAINED.md to DOCX
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_docx(md_file, docx_file):
    """Parse markdown and create DOCX"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # H1 - # Title
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # H2 - ## Title
        elif line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            
        # H3 - ### Title
        elif line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            
        # H4 - #### Title
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
            
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('_' * 80)
            doc.add_paragraph()
            
        # Code block
        elif line.startswith('```'):
            # Collect code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Intense Quote'
                
        # Bullet list - starts with - or * or •
        elif re.match(r'^[\-\*\•]\s+', line):
            # Clean text
            text = re.sub(r'^[\-\*\•]\s+', '', line)
            text = clean_markdown(text)
            
            # Check indentation for nested lists
            spaces = len(line) - len(line.lstrip())
            level = spaces // 2
            
            p = doc.add_paragraph(text, style='List Bullet')
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.5 * level)
        
        # Numbered list
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            text = clean_markdown(text)
            doc.add_paragraph(text, style='List Number')
        
        # Special formatted lines (bold with emoji)
        elif line.startswith('**') and '**' in line[2:]:
            text = clean_markdown(line)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        
        # Regular paragraph
        else:
            text = clean_markdown(line)
            if text:
                doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f'✅ DOCX файлът е създаден успешно: {docx_file}')

def clean_markdown(text):
    """Remove markdown formatting"""
    # Bold **text**
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Italic *text*
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Code `text`
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Links [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

if __name__ == '__main__':
    parse_markdown_to_docx(
        'PRESENTATION_EXPLAINED.md',
        'PRESENTATION_EXPLAINED.docx'
    )
